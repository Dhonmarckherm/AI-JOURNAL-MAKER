"""
AI Journal Maker - Vercel Serverless Entry Point
Public access - no authentication required
Using simple WSGI-style handler
"""

import os
import sys
import json
import uuid
import base64
import io
import hashlib
import sqlite3
from pathlib import Path
from datetime import datetime
from typing import List, Optional

# Set up paths
current_dir = Path(__file__).parent.resolve()
project_root = current_dir.parent

# Vercel uses /tmp for writable storage
DATA_DIR = Path("/tmp/journal_data")
IMAGES_DIR = Path("/tmp/journal_data/images")
DATA_DIR.mkdir(parents=True, exist_ok=True)
IMAGES_DIR.mkdir(parents=True, exist_ok=True)

# Template paths
TEMPLATES_DIR = project_root / "journal_templates"

# Import dependencies
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import httpx

# Database path
DB_PATH = DATA_DIR / "journals.db"
DEFAULT_USER_ID = 1

# Configuration
CONFIG = {
    "llm_provider": "openrouter",
    "model": "nvidia/nemotron-nano-12b-v2-vl:free",
    "temperature": 0.7,
    "max_tokens": 2048,
    "api_key_env": "OPENROUTER_API_KEY",
    "base_url": "https://openrouter.ai/api/v1",
    "ui": {
        "title": "AI Journal Maker",
        "max_image_size_mb": 10,
        "allowed_extensions": ["jpg", "jpeg", "png", "gif", "webp"]
    }
}

# Initialize FastAPI app
app = FastAPI(title="AI Journal Maker", version="3.0.0")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database connection
_db_conn = None


def get_db():
    """Get database connection"""
    global _db_conn
    if _db_conn is None:
        _db_conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
        _db_conn.row_factory = sqlite3.Row
        init_db(_db_conn)
    return _db_conn


def init_db(conn):
    """Initialize database"""
    cursor = conn.cursor()
    cursor.execute('PRAGMA foreign_keys = OFF')
    cursor.execute('''CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        email TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL
    )''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS journals (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        date TEXT NOT NULL,
        time TEXT NOT NULL,
        notes TEXT,
        report TEXT NOT NULL,
        analysis TEXT,
        image_paths TEXT,
        image_count INTEGER DEFAULT 0
    )''')
    cursor.execute('''INSERT OR IGNORE INTO users (id, username, email, password_hash)
        VALUES (?, ?, ?, ?)''', (1, 'guest', 'guest@public.local', 'default'))
    conn.commit()


# Pydantic model
class JournalEntry(BaseModel):
    title: str
    date: str
    time: str
    notes: Optional[str] = ""
    report: str
    analysis: Optional[str] = ""
    images: Optional[List[str]] = []
    image_count: Optional[int] = 0


# Routes
@app.get("/")
async def root():
    """Serve main UI"""
    template_path = TEMPLATES_DIR / "index.html"
    with open(template_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())


@app.get("/static/app.js")
async def serve_js():
    """Serve JavaScript"""
    template_path = TEMPLATES_DIR / "app.js"
    with open(template_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read(), media_type="application/javascript")


@app.get("/api/health")
async def health():
    """Health check"""
    try:
        conn = get_db()
        cursor = conn.cursor()
        cursor.execute("SELECT 1")
        return JSONResponse(content={
            "status": "healthy",
            "db_ok": True,
            "data_dir": str(DATA_DIR),
            "templates_exist": TEMPLATES_DIR.exists()
        })
    except Exception as e:
        return JSONResponse(status_code=500, content={"status": "error", "message": str(e)})


@app.post("/api/analyze")
async def analyze_images(
    request: Request,
    images: List[UploadFile] = File(...),
    title: str = Form(...),
    date: str = Form(...),
    time: str = Form(...),
    notes: str = Form("")
):
    """Analyze images with AI"""
    if not images:
        raise HTTPException(status_code=400, detail="No images uploaded")

    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="API key not configured")

    allowed_types = ["image/jpeg", "image/jpg", "image/png", "image/gif", "image/webp"]
    max_size = 10 * 1024 * 1024

    image_data = []
    saved_images = []

    for img in images:
        if img.content_type not in allowed_types:
            raise HTTPException(status_code=400, detail=f"Invalid file type: {img.filename}")

        content = await img.read()
        if len(content) > max_size:
            raise HTTPException(status_code=400, detail="File too large")

        unique_filename = f"{uuid.uuid4()}_{img.filename}"
        img_path = IMAGES_DIR / unique_filename
        with open(img_path, "wb") as f:
            f.write(content)
        saved_images.append(unique_filename)

        base64_data = base64.b64encode(content).decode('utf-8')
        image_data.append({"base64": base64_data, "type": img.content_type})

    # Call OpenRouter API
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://ai-journal-maker.vercel.app",
        "X-Title": "AI Journal Maker"
    }

    content_messages = [{"type": "text", "text": f"""
Analyze this journal entry and create a detailed report.
Title: {title}
Date: {date}
Time: {time}
Notes: {notes}

Please describe what you see in the images and provide insights, reflections, and analysis.
"""}]

    for img in image_data:
        content_messages.append({
            "type": "image_url",
            "image_url": {"url": f"data:{img['type']};base64,{img['base64']}"}
        })

    payload = {
        "model": "nvidia/nemotron-nano-12b-v2-vl:free",
        "messages": [{"role": "user", "content": content_messages}],
        "temperature": 0.7,
        "max_tokens": 2048
    }

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://openrouter.ai/api/v1/chat/completions",
                headers=headers,
                json=payload
            )
            response.raise_for_status()
            result = response.json()
            report = result["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

    return JSONResponse(content={
        "title": title,
        "date": date,
        "time": time,
        "notes": notes,
        "report": report,
        "analysis": True,
        "images": saved_images,
        "image_count": len(saved_images)
    })


@app.get("/api/journals")
async def get_journals():
    """Get all journals"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM journals WHERE user_id = ? ORDER BY date DESC, time DESC", (DEFAULT_USER_ID,))
    entries = []
    for row in cursor.fetchall():
        entry = dict(row)
        entry['images'] = json.loads(entry['image_paths']) if entry['image_paths'] else []
        del entry['image_paths']
        entries.append(entry)
    return JSONResponse(content=entries)


@app.get("/api/journals/{journal_id}")
async def get_journal(journal_id: int):
    """Get single journal"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM journals WHERE id = ? AND user_id = ?", (journal_id, DEFAULT_USER_ID))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Not found")
    entry = dict(row)
    entry['images'] = json.loads(entry['image_paths']) if entry['image_paths'] else []
    del entry['image_paths']
    return JSONResponse(content=entry)


@app.post("/api/journals")
async def create_journal(entry: JournalEntry):
    """Create journal"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute('''INSERT INTO journals (user_id, title, date, time, notes, report, analysis, image_paths, image_count)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
        (DEFAULT_USER_ID, entry.title, entry.date, entry.time, entry.notes or '',
         entry.report, entry.analysis or '', json.dumps(entry.images or []), entry.image_count or 0))
    conn.commit()
    return JSONResponse(content={"id": cursor.lastrowid, "message": "Saved"})


@app.put("/api/journals/{journal_id}")
async def update_journal(journal_id: int, entry: JournalEntry):
    """Update journal"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM journals WHERE id = ? AND user_id = ?", (journal_id, DEFAULT_USER_ID))
    if not cursor.fetchone():
        raise HTTPException(status_code=404, detail="Not found")
    cursor.execute('''UPDATE journals SET title=?, date=?, time=?, notes=?, report=?, image_count=?, image_paths=?
        WHERE id=? AND user_id=?''',
        (entry.title, entry.date, entry.time, entry.notes or '', entry.report,
         entry.image_count or 0, json.dumps(entry.images or []), journal_id, DEFAULT_USER_ID))
    conn.commit()
    return JSONResponse(content={"message": "Updated"})


@app.delete("/api/journals/{journal_id}")
async def delete_journal(journal_id: int):
    """Delete journal"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM journals WHERE id = ? AND user_id = ?", (journal_id, DEFAULT_USER_ID))
    conn.commit()
    return JSONResponse(content={"message": "Deleted"})


@app.get("/api/journals/{journal_id}/export-word")
async def export_word(journal_id: int):
    """Export as Word"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM journals WHERE id = ? AND user_id = ?", (journal_id, DEFAULT_USER_ID))
    row = cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Not found")

    entry = dict(row)
    doc = Document()
    doc.add_heading(entry['title'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {entry['date']} at {entry['time']}").italic = True

    if entry.get('notes'):
        doc.add_heading('Notes', level=2)
        doc.add_paragraph(entry['notes'])

    images = json.loads(entry['image_paths']) if entry['image_paths'] else []
    if images:
        doc.add_heading('Images', level=2)
        img_path = IMAGES_DIR / images[0]
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(6))
            except:
                pass

    doc.add_heading('AI Generated Report', level=2)
    report_text = entry['report'].replace('**', '').replace('__', '')
    for para in report_text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = "".join(c for c in entry['title'] if c.isalnum() or c in ' -_').strip()
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_title.replace(' ', '_')}.docx"})


@app.get("/journal_data/images/{filename}")
async def serve_image(filename: str):
    """Serve image"""
    img_path = IMAGES_DIR / filename
    if not img_path.exists():
        raise HTTPException(status_code=404, detail="Not found")
    return StreamingResponse(open(img_path, "rb"), media_type="image/jpeg")


# Vercel handler - ASGI compliant
handler = app
