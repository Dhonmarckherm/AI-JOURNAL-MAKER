"""
AI Journal Maker - Production FastAPI Application
Deployed on Railway with PostgreSQL and Cloudinary
"""

import os
import json
import uuid
import base64
import secrets
import io
import cloudinary
import cloudinary.uploader
from pathlib import Path
from datetime import datetime, timedelta
from typing import List, Optional
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

load_dotenv()

from journal_ai import JournalAIAnalyzer
from journal_db import JournalDatabase

# Load configuration
CONFIG_PATH = Path(__file__).parent / "journal_config.json"
with open(CONFIG_PATH) as f:
    CONFIG = json.load(f)

# Initialize FastAPI app
app = FastAPI(
    title="AI Journal Maker",
    description="AI-powered journal creation system with image analysis",
    version="2.0.0"
)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Cloudinary for production image storage
CLOUDINARY_URL = os.getenv("CLOUDINARY_URL")
if CLOUDINARY_URL:
    cloudinary.config(cloud_url=CLOUDINARY_URL)
    USE_CLOUD_STORAGE = True
else:
    # Local storage for development
    DATA_DIR = Path(CONFIG.get('storage', {}).get('data_dir', './journal_data'))
    IMAGES_DIR = Path(CONFIG.get('storage', {}).get('images_dir', './journal_data/images'))
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    USE_CLOUD_STORAGE = False

# Initialize database and AI analyzer
db = JournalDatabase()
ai_analyzer = JournalAIAnalyzer(CONFIG)

# Mount static files for local development
if not USE_CLOUD_STORAGE:
    app.mount("/journal_data", StaticFiles(directory=str(DATA_DIR)), name="journal_data")

# Session management
sessions = {}

def create_session(user_id: int, username: str) -> str:
    """Create a new session"""
    session_token = secrets.token_urlsafe(32)
    sessions[session_token] = {
        "user_id": user_id,
        "username": username,
        "expires_at": datetime.now() + timedelta(days=7)
    }
    return session_token

def get_session(request: Request) -> Optional[dict]:
    """Get current session from cookie"""
    session_token = request.cookies.get("session_token")
    if not session_token or session_token not in sessions:
        return None
    
    session = sessions[session_token]
    if session["expires_at"] < datetime.now():
        del sessions[session_token]
        return None
    
    return session

# Pydantic models
class JournalEntry(BaseModel):
    title: str
    date: str
    time: str
    notes: Optional[str] = ""
    report: str
    analysis: Optional[str] = ""
    images: Optional[List[str]] = []
    image_count: Optional[int] = 0

class UserRegister(BaseModel):
    username: str
    email: str
    password: str

class UserLogin(BaseModel):
    username: str
    password: str

# Routes
@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    """Serve the main UI or redirect to login"""
    session = get_session(request)
    if not session:
        return RedirectResponse(url="/login")
    
    template_path = Path(__file__).parent.parent / "journal_templates" / "index.html"
    with open(template_path, "r", encoding="utf-8") as f:
        content = f.read()
    content = content.replace('</title>', f'</title><script>window.currentUser="{session["username"]}"</script>')
    return HTMLResponse(content=content)

@app.get("/login", response_class=HTMLResponse)
async def login_page():
    """Serve login page"""
    template_path = Path(__file__).parent.parent / "journal_templates" / "login.html"
    with open(template_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/register", response_class=HTMLResponse)
async def register_page():
    """Serve register page"""
    template_path = Path(__file__).parent.parent / "journal_templates" / "register.html"
    with open(template_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/static/app.js")
async def serve_js():
    """Serve the frontend JavaScript"""
    template_path = Path(__file__).parent.parent / "journal_templates" / "app.js"
    with open(template_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read(), media_type="application/javascript")

# Authentication Endpoints
@app.post("/api/register")
async def register(user: UserRegister):
    """Register a new user"""
    if len(user.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    if len(user.username) < 3:
        raise HTTPException(status_code=400, detail="Username must be at least 3 characters")
    
    user_id = db.create_user(user.username, user.email, user.password)
    if not user_id:
        raise HTTPException(status_code=400, detail="Username or email already exists")
    
    return JSONResponse(content={"message": "Registration successful! Please login."})

@app.post("/api/login")
async def login(user: UserLogin, request: Request):
    """Login user"""
    user_info = db.verify_user(user.username, user.password)
    if not user_info:
        raise HTTPException(status_code=401, detail="Invalid username or password")
    
    session_token = create_session(user_info['id'], user_info['username'])
    
    response = JSONResponse(content={"message": "Login successful"})
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        max_age=60 * 60 * 24 * 7,
        samesite="lax"
    )
    return response

@app.post("/api/logout")
async def logout():
    """Logout user"""
    return JSONResponse(content={"message": "Logged out"})

@app.get("/api/me")
async def get_current_user_info(request: Request):
    """Get current user info"""
    session = get_session(request)
    if not session:
        return JSONResponse(content={"authenticated": False})
    
    return JSONResponse(content={
        "authenticated": True,
        "user_id": session["user_id"],
        "username": session["username"]
    })

# Journal Endpoints
@app.post("/api/analyze")
async def analyze_images(
    request: Request,
    images: List[UploadFile] = File(...),
    title: str = Form(...),
    date: str = Form(...),
    time: str = Form(...),
    notes: str = Form("")
):
    """Analyze uploaded images with AI"""
    session = get_session(request)
    if not session:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    if not images:
        raise HTTPException(status_code=400, detail="No images uploaded")
    
    allowed_types = ["image/jpeg", "image/jpg", "image/png", "image/gif", "image/webp"]
    max_size = CONFIG.get('ui', {}).get('max_image_size_mb', 10) * 1024 * 1024
    
    image_data = []
    saved_image_urls = []

    for img in images:
        if img.content_type not in allowed_types:
            raise HTTPException(status_code=400, detail=f"Invalid file type: {img.filename}")

        content = await img.read()

        if len(content) > max_size:
            raise HTTPException(status_code=400, detail=f"File exceeds size limit")

        # Upload to Cloudinary (production) or save locally (development)
        if USE_CLOUD_STORAGE:
            upload_result = cloudinary.uploader.upload_image(
                content,
                folder="journal_maker",
                public_id=f"{uuid.uuid4()}"
            )
            image_url = upload_result['secure_url']
            saved_image_urls.append(image_url)
            
            # Still need base64 for AI analysis
            base64_data = base64.b64encode(content).decode('utf-8')
            image_data.append({
                "base64": base64_data,
                "type": img.content_type,
                "filename": img.filename
            })
        else:
            # Local storage
            unique_filename = f"{uuid.uuid4()}_{img.filename}"
            img_path = IMAGES_DIR / unique_filename
            with open(img_path, "wb") as f:
                f.write(content)
            saved_image_urls.append(unique_filename)
            
            base64_data = base64.b64encode(content).decode('utf-8')
            image_data.append({
                "base64": base64_data,
                "type": img.content_type,
                "filename": img.filename
            })

    try:
        result = ai_analyzer.analyze_images(
            images=image_data,
            title=title,
            date=date,
            time=time,
            notes=notes
        )
        result['images'] = saved_image_urls
        return JSONResponse(content=result)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@app.get("/api/journals")
async def get_journals(request: Request):
    """Get all journal entries for current user"""
    session = get_session(request)
    if not session:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    entries = db.get_all_entries(session["user_id"])
    return JSONResponse(content=entries)

@app.get("/api/journals/{journal_id}")
async def get_journal(journal_id: int, request: Request):
    """Get a specific journal entry"""
    session = get_session(request)
    if not session:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    entry = db.get_entry(journal_id, session["user_id"])
    if not entry:
        raise HTTPException(status_code=404, detail="Journal entry not found")
    return JSONResponse(content=entry)

@app.post("/api/journals")
async def create_journal(entry: JournalEntry, request: Request):
    """Save a journal entry"""
    session = get_session(request)
    if not session:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    try:
        entry_id = db.add_entry(entry.dict(), session["user_id"])
        return JSONResponse(content={"id": entry_id, "message": "Journal saved successfully"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save journal: {str(e)}")

@app.delete("/api/journals/{journal_id}")
async def delete_journal(journal_id: int, request: Request):
    """Delete a journal entry"""
    session = get_session(request)
    if not session:
        raise HTTPException(status_code=401, detail="Not authenticated")

    entry = db.get_entry(journal_id, session["user_id"])
    if not entry:
        raise HTTPException(status_code=404, detail="Journal entry not found")

    db.delete_entry(journal_id, session["user_id"])
    return JSONResponse(content={"message": "Journal deleted successfully"})

@app.put("/api/journals/{journal_id}")
async def update_journal(journal_id: int, entry: JournalEntry, request: Request):
    """Update a journal entry"""
    session = get_session(request)
    if not session:
        raise HTTPException(status_code=401, detail="Not authenticated")

    existing_entry = db.get_entry(journal_id, session["user_id"])
    if not existing_entry:
        raise HTTPException(status_code=404, detail="Journal entry not found")

    try:
        db.update_entry(journal_id, entry.dict(), session["user_id"])
        return JSONResponse(content={"message": "Journal updated successfully"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update journal: {str(e)}")

@app.get("/api/journals/{journal_id}/export-word")
async def export_journal_word(journal_id: int, request: Request):
    """Export journal entry as Word document"""
    session = get_session(request)
    if not session:
        raise HTTPException(status_code=401, detail="Not authenticated")

    entry = db.get_entry(journal_id, session["user_id"])
    if not entry:
        raise HTTPException(status_code=404, detail="Journal entry not found")

    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(entry['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date and time
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Date: {entry['date']} at {entry['time']}")
    meta_run.italic = True
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add user notes if available
    if entry.get('notes'):
        doc.add_heading('Notes', level=2)
        doc.add_paragraph(entry['notes'])
    
    # Add main image if available
    if entry.get('images') and len(entry['images']) > 0:
        doc.add_heading('Images', level=2)
        try:
            if USE_CLOUD_STORAGE:
                # Download image from Cloudinary URL
                import httpx
                img_url = entry['images'][0]
                async with httpx.AsyncClient() as client:
                    response = await client.get(img_url)
                    if response.status_code == 200:
                        img_bytes = io.BytesIO(response.content)
                        image = Image.open(img_bytes)
                        img_path = os.path.join(DATA_DIR, f"temp_{uuid.uuid4()}.png")
                        image.save(img_path, 'PNG')
                        doc.add_picture(img_path, width=Inches(6))
                        os.remove(img_path)
            else:
                # Local image
                img_path = IMAGES_DIR / entry['images'][0]
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Image could not be loaded: {str(e)}")
    
    # Add AI report
    doc.add_heading('AI Generated Report', level=2)
    
    # Process report text (remove markdown formatting)
    report_text = entry['report']
    # Remove bold markers
    report_text = report_text.replace('**', '')
    report_text = report_text.replace('__', '')
    
    # Split into paragraphs
    paragraphs = report_text.split('\n\n')
    for para in paragraphs:
        para = para.strip()
        if para:
            # Handle bullet points
            if para.startswith('- ') or para.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(para[2:])
            else:
                doc.add_paragraph(para)
    
    # Add analysis info if available
    if entry.get('image_count', 0) > 0:
        doc.add_paragraph()
        analysis_note = doc.add_paragraph()
        analysis_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        note_run = analysis_note.add_run(f"Images Analyzed: {entry['image_count']}")
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Create filename
    safe_title = "".join(c for c in entry['title'] if c.isalnum() or c in ' -_').strip()
    filename = f"{safe_title.replace(' ', '_')}.docx"
    
    return StreamingResponse(
        doc_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return JSONResponse(content={
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "2.0.0"
    })

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8000"))
    uvicorn.run(app, host="0.0.0.0", port=port)
